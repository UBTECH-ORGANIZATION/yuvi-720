"""Builds the Hebrew RTL DOCX report of the Experis/MoE finding fixes.

Run:  backend/.venv/bin/python docs/720/experis-fixes/generate_report.py
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from PIL import Image

BASE = Path(__file__).resolve().parent
SHOTS = BASE / "after"
OUT = BASE / "Yuvilab-Spark-דוח-תיקון-ממצאים.docx"

PURPLE = RGBColor(0x5A, 0x37, 0xD8)
INK = RGBColor(0x2D, 0x2A, 0x4A)
GREY = RGBColor(0x4F, 0x56, 0x66)
GREEN = RGBColor(0x1F, 0x7A, 0x4C)


def rtl(paragraph):
    """Word needs both the paragraph bidi flag and per-run rtl marks for Hebrew."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        for tag in ("w:rtl", "w:cs"):
            el = OxmlElement(tag)
            rPr.append(el)
    return paragraph


def para(doc, text="", size=11, bold=False, color=INK, space_after=6, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = "Arial"
    p.paragraph_format.space_after = Pt(space_after)
    rtl(p)
    if align is not None:
        p.alignment = align
    return p


def labelled(doc, label, value, value_color=INK):
    p = doc.add_paragraph()
    lab = p.add_run(f"{label}: ")
    lab.font.bold = True
    lab.font.size = Pt(10.5)
    lab.font.color.rgb = PURPLE
    lab.font.name = "Arial"
    val = p.add_run(value)
    val.font.size = Pt(10.5)
    val.font.color.rgb = value_color
    val.font.name = "Arial"
    p.paragraph_format.space_after = Pt(4)
    rtl(p)
    return p


def bullets(doc, items, size=10.5):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(size)
        run.font.name = "Arial"
        run.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(2)
        rtl(p)


def add_shot(doc, filename, caption):
    path = SHOTS / filename
    if not path.exists():
        return
    # Keep tall captures from taking a whole page on their own.
    max_w, max_h = 6.1, 4.6
    with Image.open(path) as image:
        ratio = image.height / image.width
    width = min(max_w, max_h / ratio) if ratio else max_w
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = GREY
    run.font.name = "Arial"
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    rtl(cap)


def rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "D9CFFB")
    borders.append(bottom)
    pPr.append(borders)
    p.paragraph_format.space_after = Pt(10)


FINDINGS = [
    {
        "n": 1,
        "title": "מסיחים — אנימציות ללא אפשרות עצירה",
        "location": "מסך נחיתה (הערה רוחבית — מתקיים במקומות נוספים במוצר)",
        "finding": "מסיחים- יש אנימציות שרצות ללא אפשרות לעצירה. מתקיים במקומות נוספים במוצר, יש להתייחס כהערה רוחבית.",
        "cause": [
            "רק קובץ עיצוב אחד מתוך תשעה כיבד את העדפת מערכת ההפעלה prefers-reduced-motion.",
            "לא היה שום פקד במוצר עצמו שמאפשר לעצור את האנימציות — והתקן דורש פקד בממשק, לא רק הגדרת מערכת.",
            "אנימציות ה-SVG מסוג animateMotion במסך הנחיתה ממשיכות לרוץ גם כאשר CSS מכבה אנימציות, כי SMIL אינו מושפע מ-CSS.",
        ],
        "fix": [
            "נוסף פקד קבוע «עצירת אנימציות» המוצג בסרגל העליון של כל מסך במוצר, כולל מסך הנחיתה.",
            "הבחירה נשמרת בשרת (שדה reduce_motion ב-API של מצב הלומד) ולכן נשמרת גם אחרי רענון והתחברות מחדש.",
            "הפקד מסמן את מסמך ה-HTML, וכלל עיצוב גלובלי מעביר כל אנימציה ומעבר במוצר למצב סיום מיידי.",
            "בנוסף ממשיכה להיכבד העדפת מערכת ההפעלה prefers-reduced-motion בכל קבצי העיצוב.",
            "אנימציות ה-SMIL מוסרות לחלוטין מהעץ כשהמצב פעיל, שכן לא ניתן לעצור אותן באמצעות CSS.",
            "המצב מועבר גם אל תוכן ה-iframe של הפעילויות (game.html), שהוא מסמך נפרד.",
        ],
        "files": "frontend/src/a11y/MotionProvider.tsx, frontend/src/components/MotionToggle.tsx, frontend/src/styles/global.css, frontend/src/features/landing-login/AgentsDiagram.tsx, learning-agent/game.html, backend/learner_state.py",
        "verify": "הפעלת הפקד מעבירה את משך האנימציות מ-0.6 שניות ו-6 שניות ל-0.00001 שניות, ומספר אנימציות ה-SMIL יורד מ-12 ל-0. הבחירה נשמרת ב-API ונשמרת לאחר רענון מלא של הדף.",
        "shots": [
            ("01a-hub-animated.png", "לפני: תרשים הסוכנים במסך הנחיתה — נקודות המידע נעות ברציפות על הקווים ומהוות מסיח"),
            ("01b-hub-stopped.png", "אחרי לחיצה על «עצירת אנימציות»: אותו תרשים ללא כל תנועה — נקודות המידע הוסרו והתוכן נשאר קריא"),
        ],
    },
    {
        "n": 2,
        "title": "תיאור תמונה חסר",
        "location": "מסך ראשי",
        "finding": "אין לתמונות תיאור תמונה, יש להשלים.",
        "cause": [
            "תמונת הגיבור הראשית במסך הנחיתה הוגדרה עם תיאור ריק, למרות שהיא נושאת מידע.",
            "אייקוני SVG דקורטיביים לא סומנו כדקורטיביים ולכן הוקראו על ידי קורא מסך.",
            "בפעילות המתמטיקה, ציור הזווית הוצג ללא חלופה טקסטואלית.",
        ],
        "fix": [
            "נוסף תיאור תמונה מלא ומתורגם לתמונה הראשית: «שני תלמידים יושבים ליד שולחן עם טאבלט, וביניהם רובוט הלמידה יובי».",
            "אייקוני ה-SVG הדקורטיביים סומנו כמוסתרים מקוראי מסך.",
            "לציור הזווית נוספה חלופה טקסטואלית המוקראת לקורא מסך, לדוגמה «בתמונה: זווית של 90 מעלות, ישרה».",
            "לוגו המוצר נשאר ללא תיאור במכוון, מאחר שהשם מופיע כטקסט חי לצידו — כפי שנדרש כדי למנוע הכפלה.",
        ],
        "files": "frontend/src/features/landing-login/LandingLoginPage.tsx, learning-agent/game.html, locales/he.json, locales/en.json, locales/ar.json",
        "verify": "סריקת axe-core בכלל מסכי המוצר מחזירה אפס ממצאים בכלל image-alt.",
        "shots": [("02-hero-alt.png", "התמונה הראשית לצד תיאור התמונה שנוסף לה")],
    },
    {
        "n": 3,
        "title": "ניגודיות צבעים במסך התחברות תלמיד",
        "location": "מסך התחברות תלמיד",
        "finding": "ניגודיות צבעים לא חזקה מספיק.",
        "cause": [
            "צבע הטקסט המשני היה בהיר מדי ויצר יחס ניגודיות של כ-3.0:1 בלבד מול רקע לבן, במקום 4.5:1 הנדרש.",
            "הצבע הזה הפעיל את תוויות שלבי ההתקדמות, את חיווי «מחובר» ואת שורת האבטחה שסומנו בדו״ח.",
            "לוגו הטקסט «יובי» הוצג במילוי שקוף על גרדיאנט בהיר, מה שמנע ניגודיות מספקת.",
        ],
        "fix": [
            "צבעי הטקסט המשניים הוכהו כך שכל טקסט עומד לפחות ביחס 4.5:1 מול הרקע בפועל.",
            "ללוגו הטקסט נקבע צבע בסיס כהה, והגרדיאנט עצמו הוכהה.",
            "בנוסף נוסף חיווי מיקוד מקלדת גלוי וברור בכל המוצר.",
        ],
        "files": "frontend/src/styles/global.css, frontend/src/styles/theme.css, frontend/src/styles/learner-mapping.css",
        "verify": "סריקת axe-core במסך ההתחברות והמיפוי מחזירה אפס ממצאים בכלל color-contrast.",
        "shots": [("03-login-contrast.png", "מסך התחברות התלמיד לאחר התיקון — תוויות השלבים, חיווי «מחובר» ולוגו «יובי» עומדים ביחס הניגודיות הנדרש")],
    },
    {
        "n": 4,
        "title": "חיווי הגעה ליעד בגרירת הנקודה",
        "location": "תרגיל לדוגמה — מתמטיקה",
        "finding": "לא קיים חיווי על הגעה לגרירה של הנקודה בדוגמה.",
        "cause": [
            "התרגיל הציג את סוג הזווית בלבד ולא ציין בשום שלב שהמטרה הושגה.",
            "הנקודה הנגררת לא הייתה מסומנת על גבי הציור עצמו, אלא רק כידית של פס הגרירה.",
            "לפס הגרירה לא היו שם נגיש וערך מילולי, ולכן משתמש קורא מסך לא קיבל כל משוב.",
        ],
        "fix": [
            "נוסף חיווי «הגעת ליעד» ברור: התווית הופכת לירוקה ומציגה «הגעת ליעד — זווית חדה ✓».",
            "הנקודה הנגררת מוצגת כעת על גבי הציור, ומשנה צבע ומקבלת טבעת סימון כאשר מגיעים ליעד.",
            "לפס הגרירה נוספו שם נגיש ותיאור ערך מילולי, לדוגמה «70 מעלות, זווית חדה, הגעת ליעד».",
            "נוסף אזור הכרזה חי שמקריא כל שינוי, כולל ההגעה ליעד.",
            "התרגיל ניתן להפעלה מלאה גם במקשי החיצים, וההנחיה עודכנה בהתאם.",
        ],
        "files": "learning-agent/game.html",
        "verify": "הזזת הנקודה מעבר לטווח היעד (110 מעלות) מכבה את החיווי, וחזרה לטווח (70 מעלות) מדליקה אותו מחדש יחד עם הכרזה קולית.",
        "shots": [
            ("04a-angle-off-target.png", "לפני הגעה ליעד — הנקודה אינה מסומנת והתווית מציגה את סוג הזווית בלבד"),
            ("04b-angle-on-target.png", "בהגעה ליעד — הנקודה מסומנת בירוק עם טבעת סימון, והתווית מאשרת «הגעת ליעד»"),
        ],
    },
    {
        "n": 5,
        "title": "ניווט מקלדת בתרגיל המיון",
        "location": "תרגיל לדוגמה — רובוטיקה (הערה רוחבית לתרגילים דומים)",
        "finding": "נגישות וניווט מקלדת- אין אפשרות לשייך חלק לקטגוריה דרך ניווט מקלדת. רלוונטי גם לתרגילים דומים במקצועות אחרים.",
        "cause": [
            "פריטי המיון היו לחצנים תקינים, אך הסלים היו אלמנטים רגילים עם מאזין לחיצה בלבד.",
            "לכן ניתן היה לבחור פריט במקלדת, אך לא הייתה כל דרך להגיע לסל ולשייך אליו.",
        ],
        "fix": [
            "הסלים הומרו ללחצנים אמיתיים, ולכן הם נגישים בניווט Tab וניתנים להפעלה ב-Enter או ברווח.",
            "הפריט הנבחר מסומן במצב לחוץ הנקרא על ידי קורא מסך.",
            "נוסף אזור הכרזה חי שמודיע על הבחירה, על השיוך ועל מספר הפריטים שנותרו.",
            "ההערה הרוחבית טופלה: גם תרגילי הסידור והבחירה קיבלו לחצנים תקניים והכרזות מקבילות.",
        ],
        "files": "learning-agent/game.html",
        "verify": "התרגיל הושלם מקצה לקצה באמצעות מקלדת בלבד. לאחר כל שיוך הוכרז לדוגמה «מנוע שויך ל🔧 חומרה. נותרו 4 פריטים», ובסיום «כל הכבוד! מיינת הכול נכון».",
        "shots": [],
    },
    {
        "n": 6,
        "title": "שגיאות ניגודיות בדאשבורד המורה",
        "location": "דאשבורד מורה (הערה רוחבית)",
        "finding": "שגיאת ניגודיות- מופיע בעוד מקומות במוצר, יש להתייחס כהערה רוחבית. ממליץ על שימוש בתוסף WAVE.",
        "cause": [
            "סרגל הצבעים של תצוגת המורה כלל גוונים בהירים מדי: הבהיר שבהם הגיע ליחס 2.4:1 בלבד.",
            "צבעי הסטטוס (דורש תשומת לב, לחיזוק, טוב) נעו סביב 3.3:1 עד 4.0:1.",
            "לוגו הטקסט הוצג במילוי שקוף על גרדיאנט בהיר.",
        ],
        "fix": [
            "כל אסימוני הצבע של תצוגת המורה וצבעי הסטטוס הוכהו לעמידה ביחס 4.5:1 לפחות מול הרקע בפועל.",
            "הגרדיאנט המותגי של המוצר הוכהה בכל המסכים, כך שטקסט לבן על גביו עומד בתקן.",
            "התיקון הוחל גם על דאשבורד התלמיד, פורטל הלמידה, מסך יצירת הלומדה ומסך המנטורינג.",
        ],
        "files": "frontend/src/features/teacher-view/teacherView.css, frontend/src/features/student-dashboard/dashboard.css, frontend/src/features/learning-portal/portal.css, frontend/src/features/mentoring/mentoring.css, frontend/src/features/learning-create/create.css, frontend/src/styles/results.css",
        "verify": "סריקת axe-core בכל מסכי המוצר ובכל הלשוניות מחזירה אפס ממצאים בכלל color-contrast.",
        "shots": [("06-teacher-dashboard.png", "דאשבורד המורה לאחר התיקון — סרגל הצבעים, צבעי הסטטוס והלוגו עומדים ביחס הניגודיות הנדרש")],
    },
    {
        "n": 7,
        "title": "פונקציית הצ׳אט לא ניתנת לבדיקה",
        "location": "מסך צ׳אט מורה/תלמיד",
        "finding": "לא ניתן לבדוק את פונקציית הצ׳אט בתצורה הקיימת.",
        "cause": [
            "השיחות היו נתוני הדגמה קבועים בקוד, ותשובות התלמיד נבחרו באקראי מתוך רשימה.",
            "צד המורה וצד התלמיד החזיקו כל אחד מערך נתונים נפרד, ולכן הודעה שנשלחה בצד אחד לא הגיעה כלל לצד השני.",
            "דבר לא נשמר, ולכן רענון הדף מחק כל התכתבות.",
        ],
        "fix": [
            "נבנתה שכבת שמירה אמיתית בשרת עם אוסף ייעודי להתכתבויות, בדיוק לפי דפוס השמירה הקיים במוצר.",
            "נוספו נתיבי API לקריאת רשימת ההתכתבויות, לקריאת שיחה, לשליחת הודעה ולסימון כנקראה.",
            "שני הצדדים — תצוגת המורה ודאשבורד התלמיד — מחוברים כעת לאותה שיחה, כולל מונה הודעות שלא נקראו.",
            "תשובות ההדגמה האקראיות הוסרו לחלוטין.",
            "המסכים מתרעננים אוטומטית כל חמש שניות, כך שניתן לפתוח את שני הצדדים במקביל ולראות שיחה אמיתית.",
            "נוספה גם נגישות: רשימת הנמענים הומרה ללחצנים, אזור השיחה הוגדר כאזור חי, ולשדה הכתיבה נוספה תווית.",
        ],
        "files": "backend/classroom_store.py, backend/app/routes/classroom.py, backend/server.py, frontend/src/features/teacher-view/teacherViewApp.ts, frontend/src/features/student-dashboard/dashboardApp.ts",
        "verify": "הודעה שנשלחה מתצוגת המורה התקבלה בדאשבורד התלמיד, תשובת התלמיד חזרה לתצוגת המורה, וההתכתבות נשמרה בשרת ושרדה רענון. אומתו 11 בדיקות אוטומטיות בשרת.",
        "shots": [
            ("07a-teacher-chat.png", "צד המורה — שיחה אמיתית הנשמרת בשרת, עם רשימת תלמידים ומונה הודעות שלא נקראו"),
            ("07b-student-chat.png", "צד התלמיד — אותה שיחה בדיוק, כולל התשובה שנשלחה מצד המורה"),
        ],
    },
    {
        "n": 8,
        "title": "לא ניתן לשבץ אירועים ביומן",
        "location": "יומן מורה",
        "finding": "אין אפשרות לבדוק את הפונקציונאליות של המסך, שכן אין אופצייה לשבץ אירועים.",
        "cause": [
            "אירועי היומן היו רשימה קבועה בקוד, ללא אפשרות הוספה, עריכה או מחיקה.",
            "היומן היה נעול לחודש יוני 2026 ללא ניווט בין חודשים.",
            "תאי היומן והאירועים לא היו אלמנטים ניתנים להפעלה.",
        ],
        "fix": [
            "נוסף לחצן «אירוע חדש» וטופס מלא: כותרת, סוג אירוע, תאריך, שעה והערות.",
            "נתמכות כעת הוספה, עריכה ומחיקה של אירועים, והכול נשמר בשרת באוסף ייעודי.",
            "נוסף ניווט בין חודשים, והיומן נפתח בחודש הנוכחי במקום בתאריך קבוע.",
            "תאי היומן והאירועים הומרו ללחצנים בעלי תווית ברורה, לדוגמה «12 באוגוסט, אין אירועים — הוספת אירוע».",
            "היומן של התלמיד מציג כעת את האירועים שהמורה משבצת, בתצוגה לקריאה בלבד.",
        ],
        "files": "backend/classroom_store.py, backend/app/routes/classroom.py, frontend/src/features/teacher-view/teacherViewApp.ts, frontend/src/features/teacher-view/skeleton.html, frontend/src/features/teacher-view/teacherView.css, frontend/src/features/student-dashboard/dashboardApp.ts",
        "verify": "אירוע נוצר, נערך ונמחק דרך הממשק, וכל פעולה אומתה מול השרת. אירוע שנוצר בצד המורה הופיע ביומן התלמיד. טופס ריק נחסם בהודעת שגיאה מתאימה.",
        "shots": [
            ("08a-teacher-calendar.png", "יומן המורה עם ניווט חודשים ולחצן «אירוע חדש»"),
            ("08b-event-modal.png", "טופס שיבוץ אירוע חדש"),
            ("08c-student-calendar.png", "יומן התלמיד מציג את האירוע ששובץ על ידי המורה"),
        ],
    },
]


def build():
    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(0.8)
    sectPr = section._sectPr
    bidi = OxmlElement("w:bidi")
    sectPr.append(bidi)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ---- cover ----
    para(doc, "מדינת ישראל · משרד החינוך", 12, True, PURPLE, 2, WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "המינהל לתקשוב, טכנולוגיה ומערכות מידע", 10, False, GREY, 18, WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "דו״ח תיקון ממצאים", 22, True, INK, 4, WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "סביבה מתוקשבת: Yuvilab Spark (יובי)", 14, False, GREY, 22, WD_ALIGN_PARAGRAPH.CENTER)

    labelled(doc, "ספק", "יוביטק · מספר תאגיד 51400833")
    labelled(doc, "מענה לדו״ח", "דו״ח הערכה ראשון מטעם חברת אקספריס, מיום 29.07.2026")
    labelled(doc, "תאריך הדו״ח", "03.08.2026")
    labelled(doc, "היקף", "8 ממצאים — כולם תוקנו ואומתו", GREEN)
    labelled(doc, "אופן האימות", "סריקת נגישות אוטומטית (axe-core), בדיקות מקלדת ידניות, ובדיקות אוטומטיות בשרת")

    doc.add_paragraph()
    rule(doc)

    para(doc, "תקציר מנהלים", 15, True, PURPLE, 8)
    para(
        doc,
        "כל שמונת הממצאים שהועלו בדו״ח ההערכה טופלו במלואם. שישה ממצאים היו ליקויי נגישות — שלושה מהם "
        "סומנו בדו״ח כהערה רוחבית ולכן טופלו בכל מסכי המוצר ולא רק במקום שבו אותרו. שני ממצאים נוספים "
        "נבעו מכך שרכיבים הוצגו כהדגמה בלבד ולא ניתן היה לבדוק אותם; רכיבים אלו נבנו מחדש כפונקציונליות "
        "מלאה הנשמרת בשרת.",
        10.5,
        space_after=8,
    )
    para(
        doc,
        "בתום העבודה בוצעה סריקת נגישות אוטומטית על פני 16 מסכים ומצבי מסך שונים — כולל לשוניות הצ׳אט, "
        "היומן, חלון שיבוץ האירוע ושתי הפעילויות האינטראקטיביות. הסריקה מחזירה אפס ממצאים בכללי ניגודיות "
        "צבעים, תיאורי תמונה, שמות לחצנים, שמות קישורים ותוויות שדות.",
        10.5,
        space_after=8,
    )

    # ---- summary table ----
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    headers = ["#", "מיקום", "הממצא בקצרה", "סטטוס"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        rtl(p)

    short = [
        ("1", "מסך נחיתה (רוחבי)", "אנימציות ללא אפשרות עצירה"),
        ("2", "מסך ראשי", "תיאור תמונה חסר"),
        ("3", "התחברות תלמיד", "ניגודיות צבעים חלשה"),
        ("4", "תרגיל מתמטיקה", "אין חיווי הגעה ליעד"),
        ("5", "תרגיל רובוטיקה (רוחבי)", "אין ניווט מקלדת למיון"),
        ("6", "דאשבורד מורה (רוחבי)", "שגיאות ניגודיות"),
        ("7", "צ׳אט מורה/תלמיד", "לא ניתן לבדוק את הצ׳אט"),
        ("8", "יומן מורה", "אין אפשרות לשבץ אירועים"),
    ]
    for num, loc, desc in short:
        cells = table.add_row().cells
        for cell, text, color in zip(
            cells, [num, loc, desc, "תוקן ואומת"], [INK, INK, INK, GREEN]
        ):
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            run.font.name = "Arial"
            run.font.color.rgb = color
            if color == GREEN:
                run.font.bold = True
            rtl(p)

    doc.add_page_break()

    # ---- findings ----
    for f in FINDINGS:
        para(doc, f"ממצא {f['n']} — {f['title']}", 15, True, PURPLE, 6)
        labelled(doc, "מיקום", f["location"])

        p = doc.add_paragraph()
        lab = p.add_run("הממצא בדו״ח: ")
        lab.font.bold = True
        lab.font.size = Pt(10.5)
        lab.font.color.rgb = PURPLE
        lab.font.name = "Arial"
        val = p.add_run(f"«{f['finding']}»")
        val.font.size = Pt(10.5)
        val.font.italic = True
        val.font.name = "Arial"
        p.paragraph_format.space_after = Pt(8)
        rtl(p)

        para(doc, "סיבת השורש", 11.5, True, INK, 3)
        bullets(doc, f["cause"])
        para(doc, "", 6, space_after=2)

        para(doc, "התיקון שבוצע", 11.5, True, INK, 3)
        bullets(doc, f["fix"])
        para(doc, "", 6, space_after=2)

        labelled(doc, "קבצים ששונו", f["files"], GREY)
        labelled(doc, "אופן האימות", f["verify"])
        doc.add_paragraph()

        for filename, caption in f["shots"]:
            add_shot(doc, filename, caption)

        doc.add_page_break()

    # ---- verification appendix ----
    para(doc, "נספח — סיכום האימות", 15, True, PURPLE, 8)

    para(doc, "סריקת נגישות אוטומטית (axe-core 4.10)", 12, True, INK, 4)
    para(
        doc,
        "נסרקו הכללים color-contrast, image-alt, button-name, link-name ו-label. התוצאה: אפס ממצאים "
        "בכל אחד מהמסכים והמצבים הבאים.",
        10.5,
        space_after=6,
    )
    bullets(
        doc,
        [
            "מסך נחיתה והתחברות",
            "מיפוי לומד ומסך התוצאות",
            "דאשבורד תלמיד — כולל לשוניות הצ׳אט והיומן",
            "דאשבורד מורה — כולל לשוניות הצ׳אט, היומן וחלון שיבוץ אירוע",
            "שיחות מנטורינג",
            "פורטל הלמידה, מסך הלומדה ומסך יצירת לומדה",
            "פעילות מתמטיקה ופעילות רובוטיקה",
        ],
    )

    para(doc, "", 6, space_after=4)
    para(doc, "בדיקות מקלדת", 12, True, INK, 4)
    bullets(
        doc,
        [
            "תרגיל המיון ברובוטיקה הושלם מקצה לקצה במקלדת בלבד, ללא שימוש בעכבר.",
            "פס גרירת הזווית מופעל במקשי החיצים, וכל שינוי מוכרז יחד עם חיווי ההגעה ליעד.",
            "בכל המוצר קיים חיווי מיקוד מקלדת גלוי.",
        ],
    )

    para(doc, "", 6, space_after=4)
    para(doc, "בדיקות אוטומטיות ובנייה", 12, True, INK, 4)
    bullets(
        doc,
        [
            "11 בדיקות אוטומטיות בשרת עוברות במלואן, ומכסות שליחת הודעות, סימון כנקראה, הרשאות בין משתמשים, וכן יצירה, עריכה ומחיקה של אירועי יומן.",
            "בנייה מלאה של צד הלקוח עוברת ללא שגיאות ולא אזהרות טיפוסים.",
        ],
    )

    para(doc, "", 6, space_after=4)
    para(doc, "שינויים נוספים לבקשת הלקוח", 12, True, INK, 4)
    bullets(
        doc,
        [
            "בורר השפה הוסר מכלל מסכי המוצר.",
            "פקד עצירת האנימציות הועבר אל הסרגל העליון של כל מסך, כדי שלא יסתיר את אזור הצ׳אט.",
            "הגרדיאנט המותגי של המוצר הוכהה בכל המסכים לשיפור הניגודיות.",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
